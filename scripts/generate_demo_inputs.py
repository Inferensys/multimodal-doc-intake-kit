from __future__ import annotations

from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFilter, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parent.parent
DEMO_INPUT_DIR = ROOT / "demo" / "input"


def main() -> None:
    DEMO_INPUT_DIR.mkdir(parents=True, exist_ok=True)
    _build_contract_docx(DEMO_INPUT_DIR / "enterprise-renewal-contract.docx")
    _build_invoice_pdf(DEMO_INPUT_DIR / "field-operations-invoice.pdf")
    _build_manual_scan(DEMO_INPUT_DIR / "warehouse-reset-scan.png")


def _build_contract_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Enterprise Renewal and Services Addendum", level=0)
    document.add_paragraph("Counterparty: Northwind Analytics LLC")
    document.add_paragraph("Effective Date: 2026-04-01")
    document.add_paragraph("Jurisdiction: Delaware")
    document.add_paragraph(
        "Term and Termination. The renewal term is 12 months. Either party may terminate the "
        "agreement with forty-five (45) days written notice after the initial stabilization period."
    )
    document.add_paragraph(
        "Operational Review. The vendor must deliver a quarterly controls review and attach the "
        "incident trend appendix with each renewal checkpoint."
    )
    document.save(path)


def _build_invoice_pdf(path: Path) -> None:
    pdf = canvas.Canvas(str(path), pagesize=LETTER)
    width, height = LETTER

    pdf.setFont("Helvetica-Bold", 18)
    pdf.drawString(72, height - 72, "Cascade Field Services")
    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 92, "Invoice Number: INV-2048-APR")
    pdf.drawString(72, height - 108, "Invoice Date: 2026-04-11")
    pdf.drawString(72, height - 124, "Bill To: Redwood Grid Systems")
    pdf.drawString(72, height - 140, "Currency: USD")

    pdf.setStrokeColor(colors.black)
    pdf.line(72, height - 168, width - 72, height - 168)
    pdf.setFont("Helvetica-Bold", 10)
    pdf.drawString(72, height - 184, "Item")
    pdf.drawString(330, height - 184, "Hours")
    pdf.drawString(400, height - 184, "Rate")
    pdf.drawString(476, height - 184, "Amount")

    rows = [
        ("On-site diagnostics for line 12B", "6", "$185.00", "$1,110.00"),
        ("Emergency relay replacement", "14", "$240.00", "$3,360.00"),
        ("After-hours commissioning window", "18", "$310.00", "$5,580.00"),
        ("Protective equipment surcharge", "-", "-", "$1,795.00"),
        ("Regional travel and standby", "-", "-", "$999.50"),
    ]
    y = height - 206
    pdf.setFont("Helvetica", 10)
    for item, hours, rate, amount in rows:
        pdf.drawString(72, y, item)
        pdf.drawString(336, y, hours)
        pdf.drawString(404, y, rate)
        pdf.drawRightString(width - 72, y, amount)
        y -= 18

    pdf.line(72, y - 4, width - 72, y - 4)
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawRightString(width - 72, y - 24, "Total Due: $12,844.50")
    pdf.setFont("Helvetica", 10)
    pdf.drawString(72, y - 48, "Payment Terms: Net 15")
    pdf.drawString(72, y - 64, "Remit within 15 calendar days to avoid a scheduling hold.")
    pdf.showPage()
    pdf.save()


def _build_manual_scan(path: Path) -> None:
    image = Image.new("RGB", (1440, 1080), color=(248, 244, 236))
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.load_default()
    body_font = ImageFont.load_default()

    draw.rectangle((70, 70, 1370, 980), outline=(90, 76, 54), width=3)
    draw.text((110, 120), "Warehouse Safety Reset Procedure", fill=(32, 28, 20), font=title_font)
    draw.text((110, 170), "Revision: r7", fill=(32, 28, 20), font=body_font)
    draw.text((110, 220), "1. Lock out conveyor power before opening the relay bay.", fill=(32, 28, 20), font=body_font)
    draw.text((110, 260), "2. Verify pressure bleed-off and sign the maintenance board.", fill=(32, 28, 20), font=body_font)
    draw.text((110, 300), "3. Reset the thermal breaker and record amperage deltas.", fill=(32, 28, 20), font=body_font)
    draw.text((110, 340), "4. Quarterly review required by the owning service team.", fill=(32, 28, 20), font=body_font)
    draw.text((110, 380), "[x] Supervisor sign-off", fill=(32, 28, 20), font=body_font)
    draw.text((110, 420), "[ ] Compliance sign-off", fill=(32, 28, 20), font=body_font)
    draw.text((110, 520), "Field note: relay cabinet 12B shows a recurring heat spike after 40 minutes.", fill=(48, 38, 28), font=body_font)

    image = image.filter(ImageFilter.GaussianBlur(radius=0.3))
    image.save(path)


if __name__ == "__main__":
    main()
